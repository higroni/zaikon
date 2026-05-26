"""Documents module service implementation."""

from functools import lru_cache
from pathlib import Path
import re
from urllib.parse import unquote, urlparse

import fitz
from docx import Document

from zaikon.core.schemas import ModuleHealth
from zaikon.modules.documents.schemas import (
    ClassifyDocumentRequest,
    ClassifyDocumentResponse,
    ExtractedDocument,
    ExtractTextRequest,
    ExtractTextResponse,
)

_CYRILLIC_TO_LATIN = str.maketrans(
    {
        "\u0410": "A",
        "\u0411": "B",
        "\u0412": "V",
        "\u0413": "G",
        "\u0414": "D",
        "\u0402": "Dj",
        "\u0415": "E",
        "\u0416": "Z",
        "\u0417": "Z",
        "\u0418": "I",
        "\u0408": "J",
        "\u041a": "K",
        "\u041b": "L",
        "\u0409": "Lj",
        "\u041c": "M",
        "\u041d": "N",
        "\u040a": "Nj",
        "\u041e": "O",
        "\u041f": "P",
        "\u0420": "R",
        "\u0421": "S",
        "\u0422": "T",
        "\u040b": "C",
        "\u0423": "U",
        "\u0424": "F",
        "\u0425": "H",
        "\u0426": "C",
        "\u0427": "C",
        "\u040f": "Dz",
        "\u0428": "S",
        "\u0430": "a",
        "\u0431": "b",
        "\u0432": "v",
        "\u0433": "g",
        "\u0434": "d",
        "\u0452": "dj",
        "\u0435": "e",
        "\u0436": "z",
        "\u0437": "z",
        "\u0438": "i",
        "\u0458": "j",
        "\u043a": "k",
        "\u043b": "l",
        "\u0459": "lj",
        "\u043c": "m",
        "\u043d": "n",
        "\u045a": "nj",
        "\u043e": "o",
        "\u043f": "p",
        "\u0440": "r",
        "\u0441": "s",
        "\u0442": "t",
        "\u045b": "c",
        "\u0443": "u",
        "\u0444": "f",
        "\u0445": "h",
        "\u0446": "c",
        "\u0447": "c",
        "\u045f": "dz",
        "\u0448": "s",
    }
)
_LATIN_FOLD = str.maketrans(
    {
        "\u010d": "c",
        "\u0107": "c",
        "\u0161": "s",
        "\u0111": "dj",
        "\u017e": "z",
        "\u010c": "c",
        "\u0106": "c",
        "\u0160": "s",
        "\u0110": "dj",
        "\u017d": "z",
    }
)
_DOCUMENT_TYPE_KEYWORDS = (
    ("law", ("zakon",)),
    ("regulation", ("uredba", "uredbu")),
    ("rulebook", ("pravilnik",)),
    ("order", ("naredba", "naredbu")),
    ("strategy", ("strategija", "strategiju")),
)
_BODY_REFERENCE_PATTERNS = (
    ("law", (r"\bovim\s+zakonom\b",)),
    ("regulation", (r"\bovom\s+uredbom\b",)),
    ("rulebook", (r"\bovim\s+pravilnikom\b",)),
    ("order", (r"\bovom\s+naredbom\b",)),
    ("strategy", (r"\bstrategija\s+razvoja\b",)),
)
_SERBIAN_MONTHS = {
    "januar": "01",
    "januara": "01",
    "februar": "02",
    "februara": "02",
    "mart": "03",
    "marta": "03",
    "april": "04",
    "aprila": "04",
    "maj": "05",
    "maja": "05",
    "jun": "06",
    "juna": "06",
    "jul": "07",
    "jula": "07",
    "avgust": "08",
    "avgusta": "08",
    "septembar": "09",
    "septembra": "09",
    "oktobar": "10",
    "oktobra": "10",
    "novembar": "11",
    "novembra": "11",
    "decembar": "12",
    "decembra": "12",
}
_OFFICIAL_GAZETTE_NUMBER_RE = re.compile(
    r"sluzbeni\s+glasnik\s+(?:rs|republike\s+srbije)\s+"
    r"(?:broj|br)\s+(?P<number>\d+(?:/\d+)?)"
)
_OFFICIAL_GAZETTE_DATE_RE = re.compile(
    r"(?:broj|br)\s+(?P<number>\d+(?:/\d+)?)\s+od\s+"
    r"(?P<day>\d{1,2})\s+(?P<month>[a-z]+)\s+(?P<year>\d{4})"
)


def _normalize_for_classification(value: str) -> str:
    normalized = value.translate(_CYRILLIC_TO_LATIN).translate(_LATIN_FOLD).lower()
    return re.sub(r"[^a-z0-9]+", " ", normalized)


def path_from_uri(source_uri: str) -> Path:
    if source_uri.startswith("file://"):
        parsed = urlparse(source_uri)
        path = unquote(parsed.path)
        if re.match(r"^/[A-Za-z]:/", path):
            path = path[1:]
        if parsed.netloc:
            path = f"//{parsed.netloc}{path}"
        return Path(path)
    return Path(source_uri)


class DocumentService:
    """In-process Documents service following the public module contract."""

    def health(self) -> ModuleHealth:
        return ModuleHealth(module_name="documents")

    def extract_text(self, request: ExtractTextRequest) -> ExtractTextResponse:
        path = path_from_uri(request.source_uri)
        file_type = request.file_type.lower()

        if file_type == "txt":
            content_text, metadata = self._extract_txt(path)
        elif file_type == "pdf":
            content_text, metadata = self._extract_pdf(path)
        elif file_type == "docx":
            content_text, metadata = self._extract_docx(path)
        else:
            raise ValueError(f"Unsupported extraction file_type: {request.file_type}")

        document = ExtractedDocument(
            source_uri=request.source_uri,
            filename=request.filename,
            file_type=file_type,
            content_text=content_text,
            language_code=request.language_code,
            metadata={"character_count": len(content_text), **metadata},
        )
        return ExtractTextResponse(document=document)

    def _extract_txt(self, path: Path) -> tuple[str, dict]:
        return path.read_text(encoding="utf-8"), {"page_count": None}

    def _extract_pdf(self, path: Path) -> tuple[str, dict]:
        pages = []
        with fitz.open(path) as pdf_document:
            for page in pdf_document:
                pages.append(page.get_text("text").strip())
            page_count = pdf_document.page_count
        return "\n\n".join(page for page in pages if page), {"page_count": page_count}

    def _extract_docx(self, path: Path) -> tuple[str, dict]:
        document = Document(path)
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        return "\n".join(paragraphs), {"paragraph_count": len(paragraphs)}

    def classify_document(
        self, request: ClassifyDocumentRequest
    ) -> ClassifyDocumentResponse:
        lines = request.content_text.splitlines()
        normalized_all_lines = [
            _normalize_for_classification(line).strip() for line in lines[:80]
        ]
        first_article_index = next(
            (
                index
                for index, line in enumerate(normalized_all_lines)
                if re.match(r"^(clan\s+\d+[a-z]?|\d+)\b", line)
            ),
            20,
        )
        title_lines = lines[: min(first_article_index, 20)]
        title_area = "\n".join(title_lines)
        filename = request.filename or ""
        normalized_lines = [
            _normalize_for_classification(line).strip() for line in title_lines
        ]
        title_area_normalized = _normalize_for_classification(title_area)
        metadata = self._publication_metadata(request.content_text)

        title_candidates = []
        for line_index, line in enumerate(normalized_lines):
            for document_type, keywords in _DOCUMENT_TYPE_KEYWORDS:
                if any(
                    line == keyword or line.startswith(f"{keyword} ")
                    for keyword in keywords
                ):
                    title_candidates.append((line_index, document_type))
        if title_candidates:
            _, document_type = max(title_candidates, key=lambda item: item[0])
            return ClassifyDocumentResponse(
                document_type=document_type, confidence=0.95, metadata=metadata
            )

        body_normalized = _normalize_for_classification(request.content_text)
        for document_type, patterns in _BODY_REFERENCE_PATTERNS:
            if any(
                re.search(pattern, body_normalized)
                for pattern in patterns
            ):
                return ClassifyDocumentResponse(
                    document_type=document_type, confidence=0.65, metadata=metadata
                )

        for document_type, keywords in _DOCUMENT_TYPE_KEYWORDS:
            if any(
                re.search(rf"\b{keyword}\b", title_area_normalized)
                for keyword in keywords
            ):
                return ClassifyDocumentResponse(
                    document_type=document_type, confidence=0.55, metadata=metadata
                )

        filename_normalized = _normalize_for_classification(filename)
        for document_type, keywords in _DOCUMENT_TYPE_KEYWORDS:
            if any(
                filename_normalized.startswith(keyword)
                or f" {keyword} " in f" {filename_normalized} "
                for keyword in keywords
            ):
                return ClassifyDocumentResponse(
                    document_type=document_type, confidence=0.45, metadata=metadata
                )

        return ClassifyDocumentResponse(
            document_type="unknown", confidence=0.40, metadata=metadata
        )

    def _publication_metadata(self, content_text: str) -> dict:
        normalized = _normalize_for_classification("\n".join(content_text.splitlines()[:40]))
        gazette_numbers = [
            match.group("number")
            for match in _OFFICIAL_GAZETTE_NUMBER_RE.finditer(normalized)
        ]
        publication_dates = []
        for match in _OFFICIAL_GAZETTE_DATE_RE.finditer(normalized):
            month = _SERBIAN_MONTHS.get(match.group("month"))
            if month is None:
                continue
            publication_dates.append(
                {
                    "official_gazette_number": match.group("number"),
                    "published_at": (
                        f"{match.group('year')}-{month}-{int(match.group('day')):02d}"
                    ),
                }
            )
        return {
            "official_gazette_numbers": sorted(set(gazette_numbers)),
            "publication_dates": publication_dates,
        }


@lru_cache
def get_document_service() -> DocumentService:
    return DocumentService()
