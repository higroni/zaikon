"""Report generation service."""

from functools import lru_cache
import json
from pathlib import Path
import textwrap
from uuid import UUID

from docx import Document
import fitz

from zaikon.core.config import settings
from zaikon.modules.draft_reviews.service import get_draft_review_service
from zaikon.modules.reports.schemas import (
    GenerateReportRequest,
    GenerateReportResponse,
    ReportRecord,
)


class ReportService:
    """Generates and stores review report artifacts."""

    def __init__(self, artifact_dir: Path | None = None) -> None:
        self.root = Path(artifact_dir or settings.artifact_dir) / "reports"
        self.download_dir = self.root / "downloads"
        self.root.mkdir(parents=True, exist_ok=True)
        self.download_dir.mkdir(parents=True, exist_ok=True)

    def generate_report(self, request: GenerateReportRequest) -> GenerateReportResponse:
        if request.report_format not in {"markdown", "docx", "pdf"}:
            raise ValueError("Supported report formats: markdown, docx, pdf")
        detail = get_draft_review_service().get_draft_review(request.pipeline_run_id)
        if detail is None:
            raise KeyError(f"Draft review not found: {request.pipeline_run_id}")

        content_text = self._render_markdown(detail)
        report = ReportRecord(
            pipeline_run_id=request.pipeline_run_id,
            report_format=request.report_format,
            title=f"Izvestaj - {detail.draft_review.title}",
            finding_count=len(detail.findings),
            content_text=content_text,
            metadata={
                "draft_review_status": detail.draft_review.status.value,
                "document_type": detail.draft_review.metadata.get("document_type"),
            },
        )
        if request.report_format == "docx":
            report.metadata["download_path"] = str(self._save_docx_report(report, detail))
        if request.report_format == "pdf":
            report.metadata["download_path"] = str(self._save_pdf_report(report))
        self._save_report(report)
        return GenerateReportResponse(report=report)

    def get_report(self, report_id: UUID) -> ReportRecord | None:
        path = self.root / f"{report_id}.json"
        if not path.exists():
            return None
        return ReportRecord.model_validate(json.loads(path.read_text(encoding="utf-8")))

    def list_reports(self) -> list[ReportRecord]:
        reports = [
            ReportRecord.model_validate(json.loads(path.read_text(encoding="utf-8")))
            for path in self.root.glob("*.json")
        ]
        return sorted(
            reports,
            key=lambda report: report.created_at.isoformat(),
            reverse=True,
        )

    def get_download(self, report_id: UUID) -> tuple[bytes | str, str, str] | None:
        report = self.get_report(report_id)
        if report is None:
            return None
        if report.report_format == "docx":
            path = Path(report.metadata.get("download_path", ""))
            if not path.exists():
                return None
            return (
                path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{report.report_id}.docx",
            )
        if report.report_format == "pdf":
            path = Path(report.metadata.get("download_path", ""))
            if not path.exists():
                return None
            return path.read_bytes(), "application/pdf", f"{report.report_id}.pdf"
        return report.content_text, "text/markdown; charset=utf-8", f"{report.report_id}.md"

    def _save_report(self, report: ReportRecord) -> None:
        path = self.root / f"{report.report_id}.json"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            json.dumps(report.model_dump(mode="json"), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _save_docx_report(self, report: ReportRecord, detail) -> Path:
        document = Document()
        document.add_heading(report.title, level=1)
        document.add_paragraph(f"Pipeline run: {detail.draft_review.pipeline_run_id}")
        document.add_paragraph(f"Status provere: {detail.draft_review.status.value}")
        document.add_paragraph(f"Broj nalaza: {len(detail.findings)}")
        document.add_heading("Nalazi", level=2)

        if not detail.findings:
            document.add_paragraph("Nema nalaza.")
        for index, finding in enumerate(detail.findings, start=1):
            document.add_heading(f"{index}. {finding.title}", level=3)
            document.add_paragraph(f"Tip: {finding.finding_type}")
            document.add_paragraph(f"Rizik: {finding.risk_level.value}")
            document.add_paragraph(f"Status: {finding.status.value}")
            document.add_paragraph(f"Putanja: {finding.source_path or 'n/a'}")
            document.add_paragraph(finding.explanation)
            if finding.recommendation:
                document.add_paragraph("Preporuka:")
                document.add_paragraph(finding.recommendation)
            if finding.review_note:
                document.add_paragraph("Beleška pregleda:")
                document.add_paragraph(finding.review_note)
            related_units = finding.evidence.get("related_legal_units", [])
            if related_units:
                document.add_paragraph("Povezane pravne jedinice:")
                for related in related_units:
                    document.add_paragraph(
                        f"{related.get('filename')} {related.get('path')} "
                        f"(score: {related.get('score')})",
                        style="List Bullet",
                    )
                    quote = (related.get("quote") or "").strip()
                    if quote:
                        document.add_paragraph(f"Citat: {quote}")

        path = self.download_dir / f"{report.report_id}.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
        return path

    def _save_pdf_report(self, report: ReportRecord) -> Path:
        pdf_document = fitz.open()
        margin = 54
        page_width = 595
        page_height = 842
        line_height = 14
        max_lines = int((page_height - 2 * margin) / line_height)
        lines = self._wrap_for_pdf(report.content_text)

        for page_start in range(0, len(lines), max_lines):
            page = pdf_document.new_page(width=page_width, height=page_height)
            text = "\n".join(lines[page_start : page_start + max_lines])
            page.insert_textbox(
                fitz.Rect(margin, margin, page_width - margin, page_height - margin),
                text,
                fontsize=10,
                fontname="helv",
                lineheight=1.2,
            )

        if not lines:
            pdf_document.new_page(width=page_width, height=page_height)
        path = self.download_dir / f"{report.report_id}.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        pdf_document.save(path)
        pdf_document.close()
        return path

    def _wrap_for_pdf(self, content_text: str) -> list[str]:
        lines: list[str] = []
        for line in content_text.splitlines():
            if not line:
                lines.append("")
                continue
            stripped = line.replace("`", "")
            lines.extend(textwrap.wrap(stripped, width=88) or [""])
        return lines

    def _render_markdown(self, detail) -> str:
        lines = [
            f"# Izvestaj - {detail.draft_review.title}",
            "",
            f"- Pipeline run: `{detail.draft_review.pipeline_run_id}`",
            f"- Status provere: `{detail.draft_review.status.value}`",
            f"- Broj nalaza: {len(detail.findings)}",
            "",
            "## Nalazi",
            "",
        ]
        if not detail.findings:
            lines.append("Nema nalaza.")
            return "\n".join(lines).strip() + "\n"

        for index, finding in enumerate(detail.findings, start=1):
            lines.extend(
                [
                    f"### {index}. {finding.title}",
                    "",
                    f"- Tip: `{finding.finding_type}`",
                    f"- Rizik: `{finding.risk_level.value}`",
                    f"- Status: `{finding.status.value}`",
                    f"- Putanja: `{finding.source_path or 'n/a'}`",
                    "",
                    finding.explanation,
                    "",
                ]
            )
            if finding.recommendation:
                lines.extend(["Preporuka:", "", finding.recommendation, ""])
            if finding.review_note:
                lines.extend(["Beleška pregleda:", "", finding.review_note, ""])
            related_units = finding.evidence.get("related_legal_units", [])
            if related_units:
                lines.extend(["Povezane pravne jedinice:", ""])
                for related in related_units:
                    lines.append(
                        "- "
                        f"{related.get('filename')} "
                        f"{related.get('path')} "
                        f"(score: {related.get('score')})"
                    )
                    quote = (related.get("quote") or "").strip()
                    if quote:
                        lines.append(f"  - Citat: {quote}")
                lines.append("")
        return "\n".join(lines).strip() + "\n"


@lru_cache
def get_report_service() -> ReportService:
    return ReportService()
